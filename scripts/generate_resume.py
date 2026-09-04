import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml

from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
)
from reportlab.pdfgen import canvas
import pymupdf

def set_cell_margins(cell, top=0, bottom=0, left=0, right=0):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def generate_docx(output_path: Path):
    doc = Document()
    
    # Page setup - 0.35 in margins (perfect 1-page fit)
    for section in doc.sections:
        section.top_margin = Inches(0.35)
        section.bottom_margin = Inches(0.35)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)
    
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(8.6)
    style_normal.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    
    def add_section_heading(title, is_first=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3) if is_first else Pt(9.5)
        p.paragraph_format.space_after = Pt(2.5)
        p.paragraph_format.keep_with_next = True
        
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        
        pPr = p._element.get_or_add_pPr()
        pBdr = parse_xml(
            '<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '<w:bottom w:val="single" w:sz="6" w:space="1" w:color="94A3B8"/>'
            '</w:pBdr>'
        )
        pPr.append(pBdr)
        return p

    def add_item_header(left_title, left_sub, right_date, right_loc=None, is_first_item=False):
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.columns[0].width = Inches(5.8)
        table.columns[1].width = Inches(1.7)
        
        cell_l = table.cell(0, 0)
        cell_r = table.cell(0, 1)
        set_cell_margins(cell_l, top=0, bottom=0, left=0, right=0)
        set_cell_margins(cell_r, top=0, bottom=0, left=0, right=0)
        
        p_l = cell_l.paragraphs[0]
        p_l.paragraph_format.space_before = Pt(1.5) if is_first_item else Pt(3.5)
        p_l.paragraph_format.space_after = Pt(0.5)
        run_title = p_l.add_run(left_title)
        run_title.bold = True
        run_title.font.size = Pt(8.8)
        run_title.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        
        if left_sub:
            run_sub = p_l.add_run(f" | {left_sub}")
            run_sub.font.size = Pt(8.3)
            run_sub.font.italic = True
            run_sub.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
            
        p_r = cell_r.paragraphs[0]
        p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_r.paragraph_format.space_before = Pt(1.5) if is_first_item else Pt(3.5)
        p_r.paragraph_format.space_after = Pt(0.5)
        run_date = p_r.add_run(right_date)
        run_date.bold = True
        run_date.font.size = Pt(8.3)
        run_date.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        
        if right_loc:
            p_r.add_run(f"  ({right_loc})").font.size = Pt(7.8)

    def add_bullet(text, space_after=1.2):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.08
        p.paragraph_format.left_indent = Inches(0.18)
        
        run_t = p.add_run(text)
        run_t.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
        run_t.font.size = Pt(8.3)
        return p

    # --- Header ---
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_name.paragraph_format.space_before = Pt(0)
    p_name.paragraph_format.space_after = Pt(1.0)
    run_name = p_name.add_run("TAIWO HENRY FADENI")
    run_name.bold = True
    run_name.font.size = Pt(14.5)
    run_name.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(2.0)
    run_sub = p_sub.add_run("APPLIED AI & SOFTWARE ENGINEER | LLM SYSTEMS • RAG & AGENT WORKFLOWS • DISTRIBUTED BACKENDS")
    run_sub.bold = True
    run_sub.font.size = Pt(8.6)
    run_sub.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
    
    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_contact.paragraph_format.space_before = Pt(0)
    p_contact.paragraph_format.space_after = Pt(4.0)
    run_c = p_contact.add_run("Lagos, Nigeria (Remote / Relocation)  •  +234 706 616 1980  •  hfadeni@gmail.com\n")
    run_c.font.size = Pt(8.0)
    run_c.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
    
    run_links = p_contact.add_run("Portfolio: henryfadeni.vercel.app  •  GitHub: github.com/Protagonist01  •  LinkedIn: linkedin.com/in/henry-fadeni-ai-engineer")
    run_links.font.size = Pt(8.0)
    run_links.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

    # --- Professional Summary ---
    add_section_heading("PROFESSIONAL SUMMARY", is_first=True)
    p_sum = doc.add_paragraph()
    p_sum.paragraph_format.space_before = Pt(1)
    p_sum.paragraph_format.space_after = Pt(3)
    p_sum.paragraph_format.line_spacing = 1.08
    run_sum = p_sum.add_run(
        "Applied AI & Software Engineer with an Electrical & Electronics Engineering foundation, specializing in production-grade LLM systems, "
        "evaluated RAG architectures, autonomous agent workflows, and high-throughput asynchronous backends. Experienced in building "
        "deterministic guardrails, human-in-the-loop controls, automated evaluation suites, and observable asynchronous backend pipelines. "
        "Proven track record of shipping end-to-end applications from distributed backend architecture to polished user interfaces."
    )
    run_sum.font.size = Pt(8.4)

    # --- Technical Skills ---
    add_section_heading("TECHNICAL SKILLS")
    skills = [
        ("Programming Languages: ", "Python, SQL (PostgreSQL, DuckDB, SQLite), JavaScript, MATLAB"),
        ("AI & LLM Systems: ", "LangGraph, LangChain, RAG Architecture, Vector Databases (Pinecone, ChromaDB), Embeddings, Semantic Caching, Prompt Engineering, Guardrails & Policy Gates, Structured Outputs (Pydantic), Human-in-the-Loop (HITL), Automated Evaluations (Golden Sets, Spider Benchmark)"),
        ("Backend & Distributed Systems: ", "FastAPI, WebSockets, RESTful APIs, Server-Sent Events (SSE), Celery, Redis (Pub/Sub, Caching, Sliding-Window Rate Limiting), AsyncIO, DuckDB, PostgreSQL, SQLite, Alembic, Pandas, PySpark"),
        ("Frontend & UI Engineering: ", "React, Next.js, State Management, Interactive Data Dashboards"),
        ("DevOps, Cloud & Observability: ", "AWS, Cloudflare, Docker, GitHub Actions (CI/CD), Git, Pytest (80%+ CI Coverage Gates), Prometheus, Grafana, Loki, Supabase, Vercel"),
        ("Automation: ", "n8n (self-hosted workflows, webhook orchestration, custom code nodes), event-driven pipelines, scheduled ETL, apps integration")
    ]
    for label, val in skills:
        p_sk = doc.add_paragraph()
        p_sk.paragraph_format.space_before = Pt(0)
        p_sk.paragraph_format.space_after = Pt(1.5)
        p_sk.paragraph_format.line_spacing = 1.08
        r_lbl = p_sk.add_run(label)
        r_lbl.bold = True
        r_lbl.font.size = Pt(8.2)
        r_lbl.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
        r_val = p_sk.add_run(val)
        r_val.font.size = Pt(8.2)

    # --- Applied AI & Backend Projects ---
    add_section_heading("APPLIED AI & BACKEND PROJECTS")
    
    # Project 1: RAA
    add_item_header("Retrieval-Augmented Analytics Workspace (Text-to-SQL)", "FastAPI · DuckDB · sqlglot · Redis · SSE", "2026", is_first_item=True)
    add_bullet("Natural-language analytics interface executing sandboxed, read-only analytical queries against DuckDB with streamed SSE explanations and dynamic charts.")
    add_bullet("Built a 2-stage AST validation pipeline using sqlglot (enforcing table/column verification, write-query rejection, and injection safeguards) with self-correction retry logic.")
    add_bullet("Automated evaluation harness across an 80-pair Golden Set (adapted from Spider), achieving 96% SQL validity, 74% execution accuracy, 61% failure self-correction, and ~4.2s p95 latency.")

    # Project 2: Code Review Agent
    add_item_header("Autonomous Code Review Agent (GitHub App)", "Python · FastAPI · LangGraph · Celery · Redis · Docker", "2025")
    add_bullet("Event-driven GitHub App agent that parses pull-request diffs, retrieves relevant file context, and publishes line-level inline reviews and commit statuses.")
    add_bullet("Decoupled webhook intake from model inference using Celery workers for asynchronous job queuing, with HMAC-SHA256 webhook verification and Redis sliding-window rate limiting.")
    add_bullet("Built a pluggable multi-provider LLM abstraction (OpenAI, Anthropic, Groq, Ollama) behind an 80% CI code coverage gate.")

    # Project 3: Self-Healing Monitor
    add_item_header("Self-Healing Microservices Monitor (Autonomous SRE Agent)", "LangGraph · ChromaDB · Prometheus · Postgres · React", "2026")
    add_bullet("Incident-response agent integrating Prometheus Alertmanager webhooks, LangGraph multi-step diagnosis, and ChromaDB vector runbook retrieval.")
    add_bullet("Designed a 4-condition deterministic policy gate (confidence >= 0.75, allowlisted low-risk actions, impact checks, human approval routing) preventing destructive runaway executions with PostgreSQL audit trails.")
    add_bullet("Scored 100% action and policy correctness across simulated failure scenarios with a live React operator dashboard.")

    # Project 4: Realtime Chat
    add_item_header("Distributed Realtime Chat Backend", "FastAPI · WebSockets · Redis Pub/Sub · SQLite · Docker", "2026")
    add_bullet("Multi-room WebSocket backend scaling across workers using reference-counted Redis Pub/Sub channels (one channel per active room).")
    add_bullet("JWT authentication at handshake, Redis-hash presence tracking with multi-device deduplication, and cursor-paginated message history (15 msgs/page); verified by cross-process integration tests for multi-worker delivery and connection fault isolation.")

    # --- Experience ---
    add_section_heading("EXPERIENCE")
    
    add_item_header("Freelance Applied AI & Software Engineer", "Independent Engineering & Consulting", "Jan 2025 – Present", "Remote", is_first_item=True)
    add_bullet("Delivered LLM systems, RAG workflows, agentic automation pipelines, and backend APIs across three engagements: a Series A logistics SaaS (Germany), an e-commerce operator (Nigeria), and a recruitment agency (Netherlands). Named references available on request.")
    add_bullet("Automated lead qualification end-to-end with n8n and a classification agent, tripling qualified-lead volume (~45 → ~140/week) with no added headcount.")

    add_item_header("Electrical & Automation Engineering Intern", "Promasidor Nigeria Limited", "May 2024 – Sep 2024", "Lagos, NG")
    add_bullet("Supported maintenance, diagnostic troubleshooting, and optimization of automated PLC-driven production-line systems in a high-volume FMCG manufacturing facility.")
    add_bullet("Applied structured root-cause analysis (RCA) to resolve electrical and sensor faults, cutting recurring downtime; documented 5+ automation workflows and SOPs.")

    add_item_header("Student Research Assistant", "Communication Research Group & Control Systems Lab", "2020 – 2025", "OAU, NG")
    add_bullet("Co-developed an IoT-compatible telemetry station and simulated sensor-to-microcontroller data transmission using MATLAB and Simulink.")
    add_bullet("Modeled Signal-to-Noise Ratio (SNR) across wireless configurations and contributed to smart metering prototypes for service-based tariff billing.")

    # --- Education ---
    add_section_heading("EDUCATION")
    add_item_header("B.Eng., Electrical & Electronics Engineering", "Obafemi Awolowo University", "2019 – 2025", "Ile-Ife, Nigeria", is_first_item=True)
    
    p_train = doc.add_paragraph()
    p_train.paragraph_format.space_before = Pt(1.5)
    p_train.paragraph_format.space_after = Pt(1)
    r_tr_lbl = p_train.add_run("Additional Training: ")
    r_tr_lbl.bold = True
    r_tr_lbl.font.size = Pt(8.0)
    r_tr_val = p_train.add_run("Data Engineering Track (DataCamp, 2024–2025)  •  Software & Data Engineering (Data Epic, 2025)  •  Computational Thinking for Problem Solving (UPenn, 2022)")
    r_tr_val.font.size = Pt(8.0)
    
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"Generated DOCX: {output_path}")

class NumberedCanvas(canvas.Canvas):
    def __init__(self, *args, **kwargs):
        super(NumberedCanvas, self).__init__(*args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            super(NumberedCanvas, self).showPage()
        super(NumberedCanvas, self).save()

def generate_pdf(output_path: Path):
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=letter,
        leftMargin=24,
        rightMargin=24,
        topMargin=20,
        bottomMargin=20
    )
    
    styles = getSampleStyleSheet()
    
    c_primary = colors.HexColor("#0F172A")
    c_blue = colors.HexColor("#2563EB")
    c_text = colors.HexColor("#334155")
    c_muted = colors.HexColor("#64748B")
    c_border = colors.HexColor("#94A3B8")
    
    name_style = ParagraphStyle(
        'NameStyle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=14.0,
        leading=16.0,
        alignment=1,
        textColor=c_primary,
        spaceAfter=2.0
    )
    
    title_style = ParagraphStyle(
        'TitleStyle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=8.0,
        leading=9.8,
        alignment=1,
        textColor=c_blue,
        spaceAfter=2.0
    )
    
    contact_style = ParagraphStyle(
        'ContactStyle',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=7.4,
        leading=9.0,
        alignment=1,
        textColor=c_text,
        spaceAfter=10.0
    )
    
    section_style = ParagraphStyle(
        'SectionHeading',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=8.8,
        leading=10.5,
        textColor=c_primary,
        spaceBefore=10.0,
        spaceAfter=1.0,
        keepWithNext=True
    )
    
    body_style = ParagraphStyle(
        'Body',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=7.6,
        leading=9.3,
        textColor=c_text,
        spaceAfter=1.8
    )
    
    item_title_style = ParagraphStyle(
        'ItemTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=7.7,
        leading=9.3,
        textColor=c_primary
    )
    
    item_date_style = ParagraphStyle(
        'ItemDate',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=7.5,
        leading=9.3,
        alignment=2,
        textColor=c_primary
    )
    
    bullet_style = ParagraphStyle(
        'BulletText',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=7.5,
        leading=9.2,
        textColor=c_text,
        leftIndent=9,
        firstLineIndent=-9,
        spaceAfter=1.0
    )
    
    story = []
    
    # Header
    story.append(Paragraph("TAIWO HENRY FADENI", name_style))
    story.append(Paragraph("APPLIED AI &amp; SOFTWARE ENGINEER | LLM SYSTEMS &bull; RAG &amp; AGENT WORKFLOWS &bull; DISTRIBUTED BACKENDS", title_style))
    contact_text = (
        "Lagos, Nigeria (Remote / Relocation) &nbsp;&bull;&nbsp; +234 706 616 1980 &nbsp;&bull;&nbsp; "
        '<a href="mailto:hfadeni@gmail.com" color="#2563EB">hfadeni@gmail.com</a><br/>'
        'Portfolio: <a href="https://henryfadeni.vercel.app/" color="#2563EB">henryfadeni.vercel.app</a> &nbsp;&bull;&nbsp; '
        'GitHub: <a href="https://github.com/Protagonist01" color="#2563EB">github.com/Protagonist01</a> &nbsp;&bull;&nbsp; '
        'LinkedIn: <a href="https://www.linkedin.com/in/henry-fadeni-ai-engineer/" color="#2563EB">linkedin.com/in/henry-fadeni-ai-engineer</a>'
    )
    story.append(Paragraph(contact_text, contact_style))
    
    def section_header(title, is_first=False):
        sb = 0 if is_first else 10.0
        sec_st = ParagraphStyle(
            f'Sec_{title}',
            parent=section_style,
            spaceBefore=sb
        )
        story.append(Paragraph(title, sec_st))
        story.append(HRFlowable(width="100%", thickness=0.5, color=c_border, spaceBefore=1.0, spaceAfter=3.2))
        
    def item_row(left_title, left_stack, right_date, right_loc="", is_first_item=False):
        top_pad = 0.5 if is_first_item else 3.5
        left_p = Paragraph(f"<b>{left_title}</b>" + (f" &nbsp;|&nbsp; <i><font color='#475569'>{left_stack}</font></i>" if left_stack else ""), item_title_style)
        right_str = f"<b>{right_date}</b>"
        if right_loc:
            right_str += f" &nbsp;<font color='#64748B' size='6.5'>({right_loc})</font>"
        right_p = Paragraph(right_str, item_date_style)
        t = Table([[left_p, right_p]], colWidths=[444, 120])
        t.setStyle(TableStyle([
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('LEFTPADDING', (0,0), (-1,-1), 0),
            ('RIGHTPADDING', (0,0), (-1,-1), 0),
            ('TOPPADDING', (0,0), (-1,-1), top_pad),
            ('BOTTOMPADDING', (0,0), (-1,-1), 0.5),
        ]))
        return t

    # 1. Professional Summary
    section_header("PROFESSIONAL SUMMARY", is_first=True)
    summary_p = (
        "Applied AI &amp; Software Engineer with an Electrical &amp; Electronics Engineering foundation, specializing in production-grade LLM systems, "
        "evaluated RAG architectures, autonomous agent workflows, and high-throughput asynchronous backends. Experienced in building "
        "deterministic guardrails, human-in-the-loop controls, automated evaluation suites, and observable asynchronous backend pipelines. "
        "Proven track record of shipping end-to-end applications from distributed backend architecture to polished user interfaces."
    )
    story.append(Paragraph(summary_p, body_style))
    
    # 2. Technical Skills
    section_header("TECHNICAL SKILLS")
    skills_data = [
        ("<b>Programming Languages:</b>", "Python, SQL (PostgreSQL, DuckDB, SQLite), JavaScript, MATLAB"),
        ("<b>AI &amp; LLM Systems:</b>", "LangGraph, LangChain, RAG Architecture, Vector Databases (Pinecone, ChromaDB), Embeddings, Semantic Caching, Prompt Engineering, Guardrails &amp; Policy Gates, Structured Outputs (Pydantic), Human-in-the-Loop (HITL), Automated Evaluations (Golden Sets, Spider Benchmark)"),
        ("<b>Backend &amp; Distributed Systems:</b>", "FastAPI, WebSockets, RESTful APIs, Server-Sent Events (SSE), Celery, Redis (Pub/Sub, Caching, Sliding-Window Rate Limiting), AsyncIO, DuckDB, PostgreSQL, SQLite, Alembic, Pandas, PySpark"),
        ("<b>Frontend &amp; UI Engineering:</b>", "React, Next.js, State Management, Interactive Data Dashboards"),
        ("<b>DevOps, Cloud &amp; Observability:</b>", "AWS, Cloudflare, Docker, GitHub Actions (CI/CD), Git, Pytest (80%+ CI Coverage Gates), Prometheus, Grafana, Loki, Supabase, Vercel"),
        ("<b>Automation:</b>", "n8n (self-hosted workflows, webhook orchestration, custom code nodes), event-driven pipelines, scheduled ETL, apps integration")
    ]
    for lbl, val in skills_data:
        p = Paragraph(f"{lbl} <font color='#334155'>{val}</font>", body_style)
        story.append(p)
        
    # 3. Applied AI & Backend Projects
    section_header("APPLIED AI &amp; BACKEND PROJECTS")
    
    # Project 1: RAA
    story.append(item_row("Retrieval-Augmented Analytics Workspace (Text-to-SQL)", "FastAPI · DuckDB · sqlglot · Redis · SSE", "2026", is_first_item=True))
    story.append(Paragraph("&bull; Natural-language analytics interface executing sandboxed, read-only analytical queries against DuckDB with streamed SSE explanations and dynamic charts.", bullet_style))
    story.append(Paragraph("&bull; Built a 2-stage AST validation pipeline using sqlglot (enforcing table/column verification, write-query rejection, and injection safeguards) with self-correction retry logic.", bullet_style))
    story.append(Paragraph("&bull; Automated evaluation harness across an 80-pair Golden Set (adapted from Spider), achieving <b>96% SQL validity</b>, <b>74% execution accuracy</b>, <b>61% failure self-correction</b>, and <b>~4.2s p95 latency</b>.", bullet_style))
    
    # Project 2: Code Review Agent
    story.append(item_row("Autonomous Code Review Agent (GitHub App)", "Python · FastAPI · LangGraph · Celery · Redis · Docker", "2025"))
    story.append(Paragraph("&bull; Event-driven GitHub App agent that parses pull-request diffs, retrieves relevant file context, and publishes line-level inline reviews and commit statuses.", bullet_style))
    story.append(Paragraph("&bull; Decoupled webhook intake from model inference using Celery workers for asynchronous job queuing, with HMAC-SHA256 webhook verification and Redis sliding-window rate limiting.", bullet_style))
    story.append(Paragraph("&bull; Built a pluggable multi-provider LLM abstraction (OpenAI, Anthropic, Groq, Ollama) behind an <b>80% CI code coverage gate</b>.", bullet_style))

    # Project 3: Self-Healing Monitor
    story.append(item_row("Self-Healing Microservices Monitor (Autonomous SRE Agent)", "LangGraph · ChromaDB · Prometheus · Postgres · React", "2026"))
    story.append(Paragraph("&bull; Incident-response agent integrating Prometheus Alertmanager webhooks, LangGraph multi-step diagnosis, and ChromaDB vector runbook retrieval.", bullet_style))
    story.append(Paragraph("&bull; Designed a 4-condition deterministic policy gate (confidence &gt;= 0.75, allowlisted low-risk actions, impact checks, human approval routing) preventing destructive runaway executions with PostgreSQL audit trails.", bullet_style))
    story.append(Paragraph("&bull; Scored <b>100% action and policy correctness</b> across simulated failure scenarios with a live React operator dashboard.", bullet_style))

    # Project 4: Realtime Chat
    story.append(item_row("Distributed Realtime Chat Backend", "FastAPI · WebSockets · Redis Pub/Sub · SQLite · Docker", "2026"))
    story.append(Paragraph("&bull; Multi-room WebSocket backend scaling across workers using reference-counted Redis Pub/Sub channels (one channel per active room).", bullet_style))
    story.append(Paragraph("&bull; JWT authentication at handshake, Redis-hash presence tracking with multi-device deduplication, and cursor-paginated message history (15 msgs/page); verified by cross-process integration tests for multi-worker delivery and connection fault isolation.", bullet_style))

    # 4. Experience
    section_header("EXPERIENCE")
    
    story.append(item_row("Freelance Applied AI &amp; Software Engineer", "Independent Engineering &amp; Consulting", "Jan 2025 – Present", "Remote", is_first_item=True))
    story.append(Paragraph("&bull; Delivered LLM systems, RAG workflows, agentic automation pipelines, and backend APIs across three engagements: a Series A logistics SaaS (Germany), an e-commerce operator (Nigeria), and a recruitment agency (Netherlands). Named references available on request.", bullet_style))
    story.append(Paragraph("&bull; Automated lead qualification end-to-end with n8n and a classification agent, tripling qualified-lead volume (~45 &rarr; ~140/week) with no added headcount.", bullet_style))

    story.append(item_row("Electrical &amp; Automation Engineering Intern", "Promasidor Nigeria Limited", "May 2024 – Sep 2024", "Lagos, NG"))
    story.append(Paragraph("&bull; Supported maintenance, diagnostic troubleshooting, and optimization of automated PLC-driven production-line systems in a high-volume FMCG manufacturing facility.", bullet_style))
    story.append(Paragraph("&bull; Applied structured root-cause analysis (RCA) to resolve electrical and sensor faults, cutting recurring downtime; documented 5+ automation workflows and SOPs.", bullet_style))

    story.append(item_row("Student Research Assistant", "Communication Research Group &amp; Control Systems Lab", "2020 – 2025", "OAU, NG"))
    story.append(Paragraph("&bull; Co-developed an IoT-compatible telemetry station and simulated sensor-to-microcontroller data transmission using MATLAB and Simulink.", bullet_style))
    story.append(Paragraph("&bull; Modeled Signal-to-Noise Ratio (SNR) across wireless configurations and contributed to smart metering prototypes for service-based tariff billing.", bullet_style))

    # 5. Education
    section_header("EDUCATION")
    story.append(item_row("B.Eng., Electrical &amp; Electronics Engineering", "Obafemi Awolowo University", "2019 – 2025", "Ile-Ife, Nigeria", is_first_item=True))
    training_p = (
        "<b>Additional Training:</b> Data Engineering Track (DataCamp, 2024–2025) &nbsp;&bull;&nbsp; "
        "Software &amp; Data Engineering (Data Epic, 2025) &nbsp;&bull;&nbsp; Computational Thinking for Problem Solving (UPenn, 2022)"
    )
    story.append(Paragraph(training_p, body_style))

    # Build document
    doc.build(story, canvasmaker=NumberedCanvas)
    print(f"Generated PDF: {output_path}")

def sync_resume_files(source_pdf: Path, source_docx: Path, target_dir: Path):
    target_dir.mkdir(parents=True, exist_ok=True)
    # Both filenames maintained for backwards compatibility and clarity
    for pdf_name in ["Henry-Fadeni-Software-AI-Engineer-Resume.pdf", "Henry_Fadeni_Applied_AI_Software_Engineer.pdf"]:
        target_pdf = target_dir / pdf_name
        target_pdf.write_bytes(source_pdf.read_bytes())
    for docx_name in ["Henry_Fadeni_Software_AI_Engineer_Resume.docx", "Henry_Fadeni_Applied_AI_Software_Engineer.docx"]:
        target_docx = target_dir / docx_name
        target_docx.write_bytes(source_docx.read_bytes())

def main():
    root = Path(__file__).resolve().parent.parent
    
    # 1. Assets directory
    assets_dir = root / "assets"
    assets_dir.mkdir(exist_ok=True)
    pdf_asset = assets_dir / "Henry-Fadeni-Software-AI-Engineer-Resume.pdf"
    docx_asset = assets_dir / "Henry_Fadeni_Applied_AI_Software_Engineer.docx"
    
    generate_pdf(pdf_asset)
    generate_docx(docx_asset)
    
    # Mirror both naming conventions in assets
    (assets_dir / "Henry_Fadeni_Applied_AI_Software_Engineer.pdf").write_bytes(pdf_asset.read_bytes())
    (assets_dir / "Henry_Fadeni_Software_AI_Engineer_Resume.docx").write_bytes(docx_asset.read_bytes())
    
    # 2. Public assets directory
    public_assets_dir = root / "public" / "assets"
    sync_resume_files(pdf_asset, docx_asset, public_assets_dir)
    print(f"Synchronized with public assets directory: {public_assets_dir}")

    # 3. Dist assets directory (if dist exists)
    dist_assets_dir = root / "dist" / "assets"
    if dist_assets_dir.exists():
        sync_resume_files(pdf_asset, docx_asset, dist_assets_dir)
        print(f"Synchronized with dist assets directory: {dist_assets_dir}")

    # 4. Also sync to Codex outputs folder if it exists
    codex_outputs = Path(r"C:\Users\wolas\Documents\Codex\2026-07-16\i\outputs")
    if codex_outputs.exists():
        generate_docx(codex_outputs / "Henry_Fadeni_Software_AI_Engineer_Resume_ATS.docx")
        generate_docx(codex_outputs / "Henry_Fadeni_AI_Python_Developer_Mimic.docx")
        generate_docx(codex_outputs / "Henry_Fadeni_Applied_AI_Software_Engineer.docx")
        generate_pdf(codex_outputs / "Henry-Fadeni-Software-AI-Engineer-Resume.pdf")
        generate_pdf(codex_outputs / "Henry_Fadeni_Applied_AI_Software_Engineer.pdf")
        print("Synchronized with Codex outputs directory.")

    doc = pymupdf.open(pdf_asset)
    print(f"PDF Page Count: {len(doc)}")
    if len(doc) == 1:
        print("SUCCESS: PDF fits perfectly on exactly 1 page!")
    else:
        print(f"WARNING: PDF has {len(doc)} pages. Adjusting spacing may be required.")

if __name__ == "__main__":
    main()
